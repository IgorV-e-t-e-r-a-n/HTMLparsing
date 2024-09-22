import os
from bs4 import BeautifulSoup
from docx import Document

def convert_html_to_docx(html_file, output_path):
    with open(html_file, 'r', encoding='utf-8') as file:
        html_content = file.read()
    # Parse HTML content
    soup = BeautifulSoup(html_content, 'html.parser')
    # Create a new DOCX document
    doc = Document()
    # Extract text from HTML and add it to the DOCX document
    for paragraph in soup.find_all('p'):
        # Add a page break before each paragraph except the first one
        if len(doc.paragraphs) > 0:
            doc.add_page_break()
        doc.add_paragraph(paragraph.get_text())
    
    # Save the DOCX document
    doc.save(output_path)

def convert_html_files_to_docx(input_folder, output_folder):
    # Create output folder if it doesn't exist
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    # Iterate through HTML files in the input folder
    for filename in os.listdir(input_folder):
        if filename.endswith(".html"):
            html_file = os.path.join(input_folder, filename)
            output_path = os.path.join(output_folder, filename.replace(".html", ".docx"))
            convert_html_to_docx(html_file, output_path)
            print(f"Converted {filename} to DOCX.")
        else:
            print("Not an HTML file:", filename)

# Set input and output folder paths
input_folder = r"C:/Users/vetki.DESKTOP-9K2M5Q3/OneDrive/Рабочий стол/html converter"
output_folder = r"C:/Users/vetki.DESKTOP-9K2M5Q3/OneDrive/Converted"

# Convert HTML files to DOCX
convert_html_files_to_docx(input_folder, output_folder)
